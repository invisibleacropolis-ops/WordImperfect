"""File IO routines for the WordImperfect editor."""

from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Final, Mapping

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]


class FileService:
    """Read and write documents in the supported formats.

    Beyond the textual payload this service is also responsible for round-
    tripping supplemental metadata that the controllers track in memory.
    Paragraph styling information is persisted to a side-car JSON file so the
    editor can preserve block level formatting choices even when the primary
    document format lacks native support for those semantics (e.g., ``.txt``).
    """

    _ENCODING: Final = "utf-8"
    _DOCX_SUFFIX: Final = ".docx"
    _RTF_SUFFIX: Final = ".rtf"
    _TEXT_SUFFIX: Final = ".txt"
    _STYLE_SUFFIX: Final = ".styles.json"

    def read(self, path: Path) -> str:
        """Return the textual contents of ``path``.

        The method inspects the file extension to determine the appropriate
        decoding routine. The implementation is intentionally lightweight so the
        editor can start life without external services.
        """

        suffix = path.suffix.lower()
        if suffix == self._TEXT_SUFFIX:
            return path.read_text(encoding=self._ENCODING)
        if suffix == self._RTF_SUFFIX:
            return self._decode_rtf(path.read_text(encoding=self._ENCODING))
        if suffix == self._DOCX_SUFFIX:
            if Document is None:
                msg = "python-docx is required to read .docx files"
                raise RuntimeError(msg)
            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)

        msg = f"Unsupported file format: {path.suffix}"
        raise ValueError(msg)

    def read_with_styles(self, path: Path) -> tuple[str, dict[int, dict[str, object]]]:
        """Return the textual contents and any stored paragraph style metadata.

        The metadata format is intentionally plain: a mapping of paragraph
        indices to dictionaries containing ``alignment``, ``indent``, and
        ``list_type`` entries. The controller converts the payload back into the
        richer dataclasses it exposes to the rest of the application.
        """

        text = self.read(path)
        metadata_path = self._style_metadata_path(path)
        if not metadata_path.exists():
            return text, {}

        raw_payload = json.loads(metadata_path.read_text(encoding=self._ENCODING))
        styles: dict[int, dict[str, object]] = {}
        for key, value in raw_payload.items():
            try:
                index = int(key)
            except (TypeError, ValueError):
                continue
            if isinstance(value, dict):
                styles[index] = dict(value)
        return text, styles

    def write(self, path: Path, text: str) -> None:
        """Persist ``text`` to ``path`` using the format implied by the suffix."""

        suffix = path.suffix.lower()
        if suffix == self._TEXT_SUFFIX:
            path.write_text(text, encoding=self._ENCODING)
            return
        if suffix == self._RTF_SUFFIX:
            path.write_text(self._encode_rtf(text), encoding=self._ENCODING)
            return
        if suffix == self._DOCX_SUFFIX:
            if Document is None:
                msg = "python-docx is required to write .docx files"
                raise RuntimeError(msg)
            document = Document()
            lines = text.splitlines()
            if text.endswith("\n"):
                lines.append("")
            if not lines:
                lines = [""]
            for line in lines:
                document.add_paragraph(line)
            document.save(str(path))
            return

        msg = f"Unsupported file format: {path.suffix}"
        raise ValueError(msg)

    def write_with_styles(
        self,
        path: Path,
        text: str,
        paragraph_styles: Mapping[int, Mapping[str, object]] | None,
    ) -> None:
        """Persist ``text`` and paragraph style metadata alongside ``path``.

        When ``paragraph_styles`` is empty any existing metadata side-car is
        removed so the project does not accumulate stale files when formatting
        information is cleared.
        """

        self.write(path, text)
        metadata_path = self._style_metadata_path(path)
        if paragraph_styles:
            serialisable = {
                str(index): dict(payload)
                for index, payload in paragraph_styles.items()
            }
            metadata_path.write_text(
                json.dumps(serialisable, indent=2, sort_keys=True) + "\n",
                encoding=self._ENCODING,
            )
        elif metadata_path.exists():
            metadata_path.unlink()

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    def _style_metadata_path(self, path: Path) -> Path:
        """Return the location of the JSON file storing paragraph styles."""

        return path.with_suffix(path.suffix + self._STYLE_SUFFIX)

    # ------------------------------------------------------------------
    # RTF helpers
    # ------------------------------------------------------------------
    def _encode_rtf(self, text: str) -> str:
        """Return a compact Rich Text Format representation of ``text``."""

        escaped = (
            text.replace("\\", r"\\")
            .replace("{", r"\{")
            .replace("}", r"\}")
            .replace("\t", r"\tab ")
        )
        escaped = escaped.replace("\r\n", "\n")
        escaped = escaped.replace("\n", "\\par\n")
        buffer = io.StringIO()
        buffer.write("{\\rtf1\\ansi\n")
        buffer.write(escaped)
        buffer.write("\n}")
        return buffer.getvalue()

    def _decode_rtf(self, rtf: str) -> str:
        """Best-effort conversion of a subset of RTF into plain text."""

        result: list[str] = []
        i = 0
        length = len(rtf)
        while i < length:
            char = rtf[i]
            if char == "\\":
                i += 1
                if i >= length:
                    break
                next_char = rtf[i]
                if next_char in "\\{}":
                    result.append(next_char)
                    i += 1
                    continue

                control_start = i
                while i < length and rtf[i].isalpha():
                    i += 1
                control_word = rtf[control_start:i]

                sign = 1
                if i < length and rtf[i] == "-":
                    sign = -1
                    i += 1
                number_start = i
                while i < length and rtf[i].isdigit():
                    i += 1
                if number_start != i:
                    _ = int(rtf[number_start:i]) * sign

                if i < length and rtf[i] == " ":
                    i += 1

                if control_word in {"par", "line"}:
                    result.append("\n")
                elif control_word == "tab":
                    result.append("\t")
                continue

            if char in "{}":
                i += 1
                continue

            if char in "\r\n":
                i += 1
                continue

            result.append(char)
            i += 1

        return "".join(result)
